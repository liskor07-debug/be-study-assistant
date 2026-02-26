import streamlit as st
import requests
import json
import math
import re
import time
from datetime import datetime
from xml.etree import ElementTree as ET
from bs4 import BeautifulSoup
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from io import BytesIO
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

YC_FOLDER_ID = "b1gpkc6pomiofp2jdg6v"
YC_API_KEY = "AQVN1Ps2I3v57RaWKVdxHKE9qb176s55cOfqyIiP"

st.set_page_config(page_title="Помощник по биоэквивалентности", layout="wide")
st.title("🪰 Помощник проектирования исследований биоэквивалентности")
st.markdown("### Заполните параметры слева и нажмите «Сгенерировать синопсис»")

st.sidebar.header("Общая информация")
sponsor = st.sidebar.text_input("Спонсор", value="ООО «Исследовательский Институт»")
study_number = st.sidebar.text_input("Номер исследования", value="ПЛБ-6/2025")
protocol_version = st.sidebar.text_input("Версия протокола", value="1.0")
protocol_date = st.sidebar.date_input("Дата протокола", value=datetime.today())
research_center = st.sidebar.text_input("Исследовательский центр", value="ООО «Клинический Исследовательский Центр»")
bioanalytical_lab = st.sidebar.text_input("Биоаналитическая лаборатория", value="ООО «БиоАналитика»")

st.sidebar.header("Препарат")
inn = st.sidebar.text_input("МНН", value="Palbociclib")
test_product = st.sidebar.text_input("Тестируемый", value="Палбоциклиб")
reference_product = st.sidebar.text_input("Референтный", value="Итулси")
dosage_form = st.sidebar.text_input("Лекарственная форма", value="капсулы")
dose = st.sidebar.text_input("Дозировка", value="125 мг")
manufacturer_t = st.sidebar.text_input("Производитель", value="Россия")
manufacturer_r = st.sidebar.text_input("Производитель", value="США")
registration_number_r = st.sidebar.text_input("РУ референтного", value="ЛП-№XXXXX")
auxiliary_substances_t = st.sidebar.text_input("Вспомогательные вещества", value="лактоза, крахмал, магния стеарат")
auxiliary_substances_r = st.sidebar.text_input("Вспомогательные вещества", value="лактоза, крахмал, тальк")
storage_conditions = st.sidebar.text_input("Условия хранения", value="при температуре не выше 25°C, в защищённом от света месте")

st.sidebar.header("Популяция")
regimen = st.sidebar.selectbox("Режим приёма", ["натощак", "после еды", "оба варианта"])
gender = st.sidebar.selectbox("Пол", ["оба", "только мужчины", "только женщины"])
age_min = st.sidebar.number_input("Минимальный возраст", 0, 100, 18, 1)
age_max = st.sidebar.number_input("Максимальный возраст", 0, 100, 45, 1)
bmi_min = st.sidebar.number_input("ИМТ мин", 10.0, 50.0, 18.5, 0.1)
bmi_max = st.sidebar.number_input("ИМТ макс", 10.0, 50.0, 30.0, 0.1)
weight_min_m = st.sidebar.number_input("Мин. вес мужчины", 30.0, 150.0, 55.0, 0.5)
weight_min_f = st.sidebar.number_input("Мин. вес женщины", 30.0, 150.0, 45.0, 0.5)
weight_max = st.sidebar.number_input("Макс. вес", 50.0, 200.0, 110.0, 0.5)

st.sidebar.header("Статистика")
target_power = st.sidebar.slider("Целевая мощность", 0.7, 0.95, 0.8, 0.05)
alpha = st.sidebar.number_input("Уровень значимости α", 0.01, 0.1, 0.05, 0.01)
dropout_rate = st.sidebar.slider("Отсев после рандомизации, %", 0, 40, 15, 5)
screening_fail_rate = st.sidebar.slider("Отсев на скрининге, %", 0, 50, 20, 5)

st.sidebar.header("Фармакокинетика")
use_auto_pk = st.sidebar.checkbox("Автоматически получить PK-данные", value=True)
manual_cv = st.sidebar.number_input("CVintra", 5.0, 100.0, 25.0, 0.1)
manual_thalf = st.sidebar.number_input("T½", 0.5, 200.0, 29.0, 0.5)
tmax_hours = st.sidebar.number_input("Tmax", 0.5, 48.0, 4.0, 0.5)
expected_gmr = st.sidebar.number_input("Ожидаемое GMR", 0.85, 1.15, 0.95, 0.01)

st.sidebar.header("Дизайн исследования")
design_choice = st.sidebar.selectbox(
    "Выбор дизайна", 
    ["автоматически", "2x2", "2x2x3", "2x2x4", "parallel"],
    help="Выберите автоматический подбор или укажите конкретный дизайн"
)

study_type = st.sidebar.selectbox(
    "Тип исследования", 
    ["двухфазное", "однофазное"],
    help="Двухфазное - перекрестный дизайн, однофазное - параллельные группы"
)

st.sidebar.header("RSABE")
use_rsabe = st.sidebar.checkbox(
    "Использовать RSABE", 
    value=False, 
    help="Reference-Scaled Average Bioequivalence - для препаратов с вариабельностью >30%"
)

if use_rsabe:
    regulator = st.sidebar.selectbox(
        "Регуляторные требования", 
        ["EMA (Европа)", "FDA (США)"],
        help="EMA ограничивает расширение границ для AUC, FDA разрешает для Cmax и AUC"
    )
    cv_threshold = st.sidebar.slider(
        "Порог вариабельности для RSABE", 
        20, 50, 30, 5,
        help="CV% выше которого применяется RSABE"
    )
else:
    regulator = "EMA (Европа)"
    cv_threshold = 30

st.sidebar.header("Источники")
use_pkdb = st.sidebar.checkbox("PK-DB", value=True)
use_pubmed = st.sidebar.checkbox("PubMed", value=True)
use_grls = st.sidebar.checkbox("ГРЛС", value=True)
use_drugbank = st.sidebar.checkbox("DrugBank", value=False)  
use_llm = st.sidebar.checkbox("Использовать Yandex GPT для извлечения", value=True)

run_button = st.sidebar.button("Сгенерировать синопсис", type="primary")

def call_yandex_gpt(prompt, max_tokens=2000):
    if not YC_API_KEY or not YC_FOLDER_ID:
        st.error("Не настроены ключи доступа Yandex GPT")
        return None
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {"Authorization": f"Api-Key {YC_API_KEY}", "Content-Type": "application/json"}
    data = {
        "modelUri": f"gpt://{YC_FOLDER_ID}/yandexgpt/latest",
        "completionOptions": {"stream": False, "temperature": 0.1, "maxTokens": max_tokens},
        "messages": [
            {"role": "system", "text": "Ты эксперт по клинической фармакологии и биоэквивалентности. Извлекай численные значения PK параметров из текстов. Отвечай ТОЛЬКО в формате JSON."},
            {"role": "user", "text": prompt}
        ]
    }
    for attempt in range(3):
        try:
            response = requests.post(url, headers=headers, json=data, timeout=60)
            response.raise_for_status()
            result = response.json()
            return result.get("result", {}).get("alternatives", [{}])[0].get("message", {}).get("text", "")
        except requests.exceptions.Timeout:
            if attempt < 2:
                time.sleep(2)
                continue
            st.warning("Таймаут при обращении к Yandex GPT")
        except Exception as e:
            st.warning(f"Ошибка Yandex GPT: {e}")
            if attempt < 2:
                time.sleep(2)
                continue
    return None


def extract_pk_params_from_text(text, inn, source=""):
    prompt = f"""
Ты — эксперт по клинической фармакологии и биоэквивалентности лекарственных препаратов. 
Твоя задача — найти и извлечь численные значения фармакокинетических параметров из текста 
для препарата {inn} (источник: {source}).

ВНИМАНИЕ: Верни ТОЛЬКО JSON в ответе. Никакого текста до или после JSON.

Найди следующие параметры (если не найдены — укажи null):

1. cv_intra — коэффициент внутрииндивидуальной вариации (В %!)
   Ищи как: CV, коэффициент вариации, intraindividual CV, intra-subject CV, %CV, 
   межиндивидуальная вариабельность, вариабельность, variability, coefficient of variation
   Примеры: "CV составил 25%", "коэффициент вариации 30%", "intraindividual CV = 28%"

2. t_half — период полувыведения (В ЧАСАХ!)
   Ищи как: T½, T1/2, half-life, период полувыведения, elimination half-life, t1/2, 
   терминальный период полувыведения, half life
   Примеры: "T½ = 12 часов", "период полувыведения 8.5 ч", "half-life was 29 hours"

3. cmax — максимальная концентрация (в нг/мл или мкг/л)
   Ищи как: Cmax, C max, максимальная концентрация, peak concentration, Cmax, 
   пиковая концентрация, максимальная плазменная концентрация
   Примеры: "Cmax = 150 нг/мл", "максимальная концентрация 200 mcg/L", "Cmax was 180 ng/mL"

4. auc — площадь под кривой (в нг·ч/мл или мкг·ч/л)
   Ищи как: AUC, AUC0-t, AUC0-inf, AUC0-∞, площадь под кривой, area under the curve, 
   площадь под фармакокинетической кривой
   Примеры: "AUC = 1200 нг·ч/мл", "AUC0-t = 1500", "area under curve 1800 ng·h/mL"

5. tmax — время достижения Cmax (В ЧАСАХ!)
   Ищи как: Tmax, T max, время достижения Cmax, time to Cmax, time to peak, 
   время достижения максимальной концентрации, TMAX
   Примеры: "Tmax = 4 часа", "время достижения пика 2.5 ч", "time to peak was 3 hours"

ВАЖНЫЕ ПРАВИЛА:
1. Все значения должны быть числами (float), не строками
2. Если параметр не найден — укажи null (не 0, не "N/A")
3. Конвертируй единицы: минуты → часы (раздели на 60), дни → часы (умножь на 24)
4. Если найдено несколько значений — возьми среднее или наиболее релевантное для однократной дозы
5. Ищи в любых разделах: аннотация, методы, результаты, таблицы, выводы
6. Обращай внимание на контекст — параметры должны относиться к препарату {inn}

Текст для анализа:
{text}...

ФОРМАТ ОТВЕТА (строго JSON):
{{
    "cv_intra": 25.0 или null,
    "t_half": 12.5 или null,
    "cmax": 150.0 или null,
    "auc": 1200.0 или null,
    "tmax": 4.0 или null
}}

Начинай ответ сразу с {{ и заканчивай }}. Никакого текста вне JSON.
"""

    response = call_yandex_gpt(prompt, max_tokens=1500)
    if response:
        try:
            json_match = re.search(r'\{.*?\}', response, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
                normalized = {}
                for k, v in parsed.items():
                    k_norm = k.lower().replace('_', '').replace('-', '').replace(' ', '')
                    if any(x in k_norm for x in ['cv', 'коэффициент', 'variation', 'вариаци', 'вариаб']):
                        normalized['cv_intra'] = v
                    elif any(x in k_norm for x in ['thalf', 't½', 't12', 'полувывед', 'halflife', 'half', 'halfl']):
                        normalized['t_half'] = v
                    elif any(x in k_norm for x in ['cmax', 'cmax', 'максимальн', 'peak', 'пик', 'maxcon']):
                        normalized['cmax'] = v
                    elif any(x in k_norm for x in ['auc', 'площад', 'areaunder', 'area', 'подкрив']):
                        normalized['auc'] = v
                    elif any(x in k_norm for x in ['tmax', 'tmax', 'врем', 'timeto', 'время', 'достиж']):
                        normalized['tmax'] = v
                if not normalized:
                    normalized = parsed
                res = {"cv_intra": None, "t_half": None, "cmax": None, "auc": None, "tmax": None}
                for key in res:
                    val = normalized.get(key)
                    if val is not None and val != "null" and val != "None":
                        try:
                            res[key] = float(val)
                        except:
                            if isinstance(val, str):
                                num = re.search(r'(\d+\.?\d*)', val)
                                if num:
                                    res[key] = float(num.group(1))

                return res

        except Exception as e:
            st.warning(f"Ошибка разбора JSON от Yandex GPT: {e}")
            st.warning(f"Получен ответ: {response[:200]}...")

    return {"cv_intra": None, "t_half": None, "cmax": None, "auc": None, "tmax": None}

def safe_request(url, params=None, timeout=45, retries=2, suppress_403=False):
    for attempt in range(retries + 1):
        try:
            response = requests.get(url, params=params, timeout=timeout)
            if response.status_code == 403 and suppress_403:
                return None
            response.raise_for_status()
            return response
        except requests.exceptions.Timeout:
            if attempt < retries:
                time.sleep(2)
                continue
            st.warning(f"Таймаут при запросе {url[:50]}...")
            return None
        except requests.exceptions.HTTPError as e:
            if hasattr(response, 'status_code') and response.status_code == 403 and suppress_403:
                return None
            if attempt < retries:
                time.sleep(2)
                continue
            st.warning(f"Ошибка HTTP {response.status_code if 'response' in locals() else '?'} при запросе {url[:50]}...")
            return None
        except requests.exceptions.RequestException as e:
            if attempt < retries:
                time.sleep(2)
                continue
            st.warning(f"Ошибка запроса {url[:50]}...: {e}")
            return None
    return None

def fetch_pk_data_pubmed(inn):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    collected = {"cv_intra": [], "t_half": [], "cmax": [], "auc": [], "tmax": []}
    studies = []
    try:
        query = f"({inn}[Title/Abstract]) AND (pharmacokinetics OR bioavailability OR bioequivalence OR Cmax OR AUC OR half-life OR 'coefficient of variation')"
        st.info(f"Поиск в PubMed: {query[:100]}...")
        search_resp = safe_request(f"{base_url}esearch.fcgi",
                                   params={"db": "pubmed", "term": query, "retmax": 150, "retmode": "json", "sort": "relevance"},
                                   timeout=45, retries=2)
        if search_resp and search_resp.status_code == 200:
            search_data = search_resp.json()
            pmids = search_data.get("esearchresult", {}).get("idlist", [])
            if pmids:
                st.success(f"Найдено {len(pmids)} статей")
                fetch_resp = safe_request(f"{base_url}efetch.fcgi",
                                          params={"db": "pubmed", "id": ",".join(pmids), "retmode": "xml", "rettype": "abstract"},
                                          timeout=45, retries=2)
                if fetch_resp and fetch_resp.status_code == 200:
                    root = ET.fromstring(fetch_resp.content)
                    for i, article in enumerate(root.findall(".//PubmedArticle"), 1):
                        try:
                            pmid = article.findtext(".//PMID")
                            title = article.findtext(".//ArticleTitle") or "Без названия"
                            abstract = " ".join(e.text for e in article.findall(".//AbstractText") if e.text)
                            journal = article.findtext(".//Journal/Title") or "Неизвестный журнал"
                            year = article.findtext(".//PubDate/Year") or article.findtext(".//PubDate/MedlineDate") or "Неизвестно"
                            authors = []
                            for a in article.findall(".//Author"):
                                last = a.findtext("LastName")
                                fore = a.findtext("ForeName")
                                if last and fore:
                                    authors.append(f"{last} {fore}")
                            authors_text = ", ".join(authors[:3]) + (" et al." if len(authors) > 3 else "")
                            if abstract and use_llm and len(abstract) > 100:
                                st.info(f"Анализ статьи {i}: {title[:50]}...")
                                extracted = extract_pk_params_from_text(abstract, inn, f"PubMed PMID:{pmid}")
                                study_info = {
                                    "pmid": pmid, "title": title, "journal": journal, "year": year,
                                    "authors": authors_text,
                                    "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else None,
                                    "source": "PubMed", "extracted_params": extracted,
                                    "abstract": abstract[:300] + "..." if len(abstract) > 300 else abstract
                                }
                                for k in collected:
                                    if extracted.get(k):
                                        collected[k].append(extracted[k])
                                        st.write(f"Найден {k}: {extracted[k]}")
                                studies.append(study_info)
                        except Exception as e:
                            st.warning(f"Ошибка обработки статьи: {e}")
            else:
                st.info("В PubMed статей не найдено.")
    except Exception as e:
        st.warning(f"Ошибка при запросе PubMed: {e}")

    result = {}
    for k in collected:
        result[k] = (sum(collected[k]) / len(collected[k])) if collected[k] else None
    result["studies"] = studies
    result["source"] = "PubMed"
    found_params = [f"{k.upper()}={v:.2f}" for k, v in result.items()
                    if k not in ["studies", "source"] and v is not None]
    if found_params:
        st.write(f" **Параметры из PubMed:** {', '.join(found_params)}")
    elif studies:
        st.info("Статьи найдены, но PK-параметры не извлечены из аннотаций")
    return result

def fetch_pk_data_grls(inn):
    collected = {"cv_intra": [], "t_half": [], "cmax": [], "auc": [], "tmax": []}
    studies = []
    try:
        url = f"https://grls.rosminzdrav.ru/grls.aspx?p=1&t={inn}"
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        st.info(f"Поиск в ГРЛС: {inn}")
        resp = safe_request(url, timeout=45, retries=2)
        if resp and resp.status_code == 200:
            soup = BeautifulSoup(resp.text, "html.parser")
            links = [a for a in soup.find_all("a", href=True) if "grls.aspx?RegNumber=" in a["href"]]
            if links:
                st.success(f"Найдено {len(links)} препаратов")
                for i, link in enumerate(links[:3], 1):
                    try:
                        reg = re.search(r'RegNumber=([^&]+)', link["href"])
                        if not reg:
                            continue
                        reg_number = reg.group(1)
                        drug_url = "https://grls.rosminzdrav.ru/" + link["href"]
                        name = link.text.strip()
                        st.info(f"Анализ препарата {i}: {name}")
                        drug_resp = safe_request(drug_url, timeout=45, retries=1)
                        if drug_resp and drug_resp.status_code == 200:
                            drug_soup = BeautifulSoup(drug_resp.text, "html.parser")
                            instr_url = None
                            for a in drug_soup.find_all("a", href=True):
                                href = a["href"].lower()
                                if "instruction" in href or "инструкция" in a.text.lower():
                                    instr_url = a["href"] if a["href"].startswith("http") else "https://grls.rosminzdrav.ru/" + a["href"]
                                    break
                            text = drug_soup.get_text()
                            if instr_url:
                                instr_resp = safe_request(instr_url, timeout=45, retries=1)
                                if instr_resp and instr_resp.status_code == 200:
                                    instr_soup = BeautifulSoup(instr_resp.text, "html.parser")
                                    text = instr_soup.get_text()
                                    st.success("Найдена инструкция")
                            if use_llm and len(text) > 200:
                                extracted = extract_pk_params_from_text(text, inn, f"ГРЛС {name}")
                                study_info = {
                                    "name": name, "reg_number": reg_number, "url": drug_url,
                                    "instruction_url": instr_url, "source": "ГРЛС", "extracted_params": extracted
                                }
                                for k in collected:
                                    if extracted.get(k):
                                        collected[k].append(extracted[k])
                                        st.write(f"Найден {k}: {extracted[k]}")
                                studies.append(study_info)
                    except Exception as e:
                        st.warning(f"Ошибка при парсинге страницы ГРЛС: {e}")
            else:
                st.info("В ГРЛС препаратов не найдено.")
    except Exception as e:
        st.warning(f"Ошибка при запросе ГРЛС: {e}")

    result = {}
    for k in collected:
        result[k] = (sum(collected[k]) / len(collected[k])) if collected[k] else None
    result["studies"] = studies
    result["source"] = "ГРЛС"
    found_params = [f"{k.upper()}={v:.2f}" for k, v in result.items()
                    if k not in ["studies", "source"] and v is not None]
    if found_params:
        st.write(f" **Параметры из ГРЛС:** {', '.join(found_params)}")
    elif studies:
        st.info("Препараты найдены, но PK-параметры не извлечены")
    return result

def fetch_pk_data_drugbank(inn):
    collected = {"cv_intra": [], "t_half": [], "cmax": [], "auc": [], "tmax": []}
    studies = []
    st.info(f"Поиск в DrugBank: {inn}")
    try:
        search_url = f"https://go.drugbank.com/unearth/q?query={inn}&searcher=drugs"
        headers = {"User-Agent": "Mozilla/5.0"}
        resp = safe_request(search_url, timeout=45, retries=1, suppress_403=True)
        if resp and resp.status_code == 200:
            soup = BeautifulSoup(resp.text, "html.parser")
            for link in soup.find_all("a", href=True):
                if "/drugs/" in link["href"]:
                    drug_url = "https://go.drugbank.com" + link["href"]
                    name = link.text.strip()
                    try:
                        drug_resp = safe_request(drug_url, timeout=45, retries=1, suppress_403=True)
                        if drug_resp and drug_resp.status_code == 200:
                            drug_soup = BeautifulSoup(drug_resp.text, "html.parser")
                            pk_text = ""
                            for dt in drug_soup.find_all("dt"):
                                if "pharmacokinetic" in dt.text.lower():
                                    dd = dt.find_next("dd")
                                    if dd:
                                        pk_text = dd.get_text()
                                        break
                            if pk_text and use_llm:
                                extracted = extract_pk_params_from_text(pk_text, inn, "DrugBank")
                                study_info = {"name": name, "url": drug_url, "source": "DrugBank", "extracted_params": extracted}
                                for k in collected:
                                    if extracted.get(k):
                                        collected[k].append(extracted[k])
                                studies.append(study_info)
                                break
                    except:
                        continue
    except Exception as e:
        pass

    result = {}
    for k in collected:
        result[k] = (sum(collected[k]) / len(collected[k])) if collected[k] else None
    result["studies"] = studies
    result["source"] = "DrugBank"
    if studies:
        st.success(f"Найдено {len(studies)} записей в DrugBank")
        found_params = [f"{k.upper()}={v:.2f}" for k, v in result.items()
                        if k not in ["studies", "source"] and v is not None]
        if found_params:
            st.write(f" **Параметры из DrugBank:** {', '.join(found_params)}")
        elif studies:
            st.info("Препарат найден, но PK-параметры не извлечены")
    else:
        st.info("В DrugBank данных не найдено.")
    return result


def fetch_pk_data_pkdb(inn):
    BASE_URL = "https://pk-db.com/api/v1"
    ENDPOINTS = {
        "statistics": f"{BASE_URL}/statistics/substances/",
        "studies": f"{BASE_URL}/studies/",
        "pkdata_studies": f"{BASE_URL}/pkdata/studies/",
        "pkdata_data": f"{BASE_URL}/pkdata/data/",
        "references": f"{BASE_URL}/references/",
    }

    params = {
        "cv_intra": [],
        "t_half": [],
        "cmax": [],
        "auc": [],
        "studies": []
    }

    def safe_get(url, params=None, timeout=10):
        try:
            response = requests.get(url, params=params, timeout=timeout)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            st.warning(f"Запрос к {url} не удался: {e}")
            return None
        except json.JSONDecodeError:
            st.warning(f"Не удалось распарсить JSON из {url}")
            return None

    def extract_results(data):
        if data is None:
            return []
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            return data.get("results", [])
        return []

    stats_url = ENDPOINTS["statistics"]
    stats_data = safe_get(stats_url, params={"search": inn})

    if stats_data:
        results = extract_results(stats_data)
        for item in results:
            study_id = item.get("id") or item.get("study_id")
            if study_id and study_id not in params["studies"]:
                params["studies"].append(study_id)

            pk_params = item.get("pk_params", {})

            for key in ["cv_intra", "t_half", "cmax", "auc"]:
                val = item.get(key) or pk_params.get(key)
                if val is not None:
                    try:
                        params[key].append(float(val))
                    except (ValueError, TypeError):
                        pass

            if "cv" in item and item["cv"] is not None:
                try:
                    params["cv_intra"].append(float(item["cv"]))
                except (ValueError, TypeError):
                    pass
            if "half_life" in item and item["half_life"] is not None:
                try:
                    params["t_half"].append(float(item["half_life"]))
                except (ValueError, TypeError):
                    pass

    pk_studies_url = ENDPOINTS["pkdata_studies"]
    pk_studies_data = safe_get(pk_studies_url, params={"substance__icontains": inn})

    if pk_studies_data:
        results = extract_results(pk_studies_data)
        for study in results:
            study_id = study.get("id")
            if study_id and study_id not in params["studies"]:
                params["studies"].append(study_id)

            pk = study.get("pk_summary", {}) or study.get("parameters", {}) or study
            for key in ["cv_intra", "t_half", "cmax", "auc"]:
                val = pk.get(key)
                if val is not None:
                    try:
                        params[key].append(float(val))
                    except (ValueError, TypeError):
                        pass

    studies_url = ENDPOINTS["studies"]
    studies_data = safe_get(studies_url, params={"search": inn})

    if studies_data:
        results = extract_results(studies_data)
        for study in results:
            study_id = study.get("id")
            if study_id and study_id not in params["studies"]:
                params["studies"].append(study_id)

            metadata = study.get("metadata", {}) or study.get("pk_data", {}) or study
            for key in ["cv_intra", "t_half", "cmax", "auc"]:
                val = metadata.get(key)
                if val is not None:
                    try:
                        params[key].append(float(val))
                    except (ValueError, TypeError):
                        pass

    pk_data_url = ENDPOINTS["pkdata_data"]
    pk_data_response = safe_get(pk_data_url, params={"substance": inn})

    if pk_data_response:
        results = extract_results(pk_data_response)
        for entry in results:
            stats = entry.get("statistics", {})
            for key in ["cv_intra", "t_half", "cmax", "auc"]:
                val = stats.get(key) or stats.get(f"{key}_mean") or stats.get(f"{key}_median")
                if val is not None:
                    try:
                        params[key].append(float(val))
                    except (ValueError, TypeError):
                        pass

    result = {}
    for key in ["cv_intra", "t_half", "cmax", "auc"]:
        if params[key]:
            valid_vals = [v for v in params[key] if isinstance(v, (int, float)) and v > 0]
            if valid_vals:
                result[key] = sum(valid_vals) / len(valid_vals)
            else:
                result[key] = None
        else:
            result[key] = None

    result["studies"] = params["studies"]
    result["raw_counts"] = {k: len([v for v in params[k] if v is not None]) for k in
                            ["cv_intra", "t_half", "cmax", "auc"]}
    print(f"pkdb results: {result}")
    return result


def fetch_pk_data_all(inn, use_pubmed=True, use_grls=True, use_drugbank=True, use_pkdb=True):
    param_keys = ["cv_intra", "t_half", "cmax", "auc", "tmax"]
    aggregated = {k: [] for k in param_keys}
    aggregated["studies"] = []
    aggregated["sources_used"] = []

    progress = st.progress(0)
    status = st.empty()

    if use_pkdb:
        status.text("Получение данных из PK-DB...")
        try:
            data = fetch_pk_data_pkdb(inn)
            aggregated["sources_used"].append("PK-DB")
            for k in param_keys:
                if k in data and data[k] is not None:
                    aggregated[k].append(data[k])
            if data.get("studies"):
                aggregated["studies"].extend(data["studies"])
        except Exception as e:
            st.warning(f"Ошибка PK-DB: {e}")
        progress.progress(25)

    if use_pubmed:
        status.text("Получение данных из PubMed...")
        data = fetch_pk_data_pubmed(inn)
        aggregated["sources_used"].append("PubMed")
        for k in param_keys:
            if k in data and data[k] is not None:
                aggregated[k].append(data[k])
        if data.get("studies"):
            aggregated["studies"].extend(data["studies"])
        progress.progress(50)

    if use_grls:
        status.text("Получение данных из ГРЛС...")
        data = fetch_pk_data_grls(inn)
        aggregated["sources_used"].append("ГРЛС")
        for k in param_keys:
            if k in data and data[k] is not None:
                aggregated[k].append(data[k])
        if data.get("studies"):
            aggregated["studies"].extend(data["studies"])
        progress.progress(75)

    if use_drugbank:
        status.text("Получение данных из DrugBank...")
        data = fetch_pk_data_drugbank(inn)
        has_data = any(data.get(k) is not None for k in param_keys if k in data)
        if data.get("studies") or has_data:
            aggregated["sources_used"].append("DrugBank")
            for k in param_keys:
                if k in data and data[k] is not None:
                    aggregated[k].append(data[k])
            if data.get("studies"):
                aggregated["studies"].extend(data["studies"])
        else:
            st.info("DrugBank не вернул данных")
        progress.progress(100)

    status.text("Обработка завершена")
    time.sleep(1)
    progress.empty()
    status.empty()
    result = {}
    for k in param_keys:
        if aggregated[k]:
            vals = aggregated[k]
            if len(vals) > 1:
                mean = sum(vals) / len(vals)
                std = (sum((x - mean) ** 2 for x in vals) / len(vals)) ** 0.5
                filtered = [x for x in vals if abs(x - mean) <= 3 * std]
                result[k] = sum(filtered) / len(filtered) if filtered else mean
            else:
                result[k] = vals[0]
        else:
            result[k] = None

    result["studies"] = aggregated["studies"]
    result["sources_used"] = aggregated["sources_used"]
    st.success("**Итоговые PK параметры:**")
    cols = st.columns(5)
    items = [
        ("CVintra", result["cv_intra"], "%"),
        ("T½", result["t_half"], "ч"),
        ("Cmax", result["cmax"], ""),
        ("AUC", result["auc"], ""),
        ("Tmax", result["tmax"], "ч")
    ]
    for col, (name, val, unit) in zip(cols, items):
        with col:
            if val is not None:
                st.metric(name, f"{val:.2f} {unit}")
            else:
                st.metric(name, "—")

    return result
def calculate_washout_period(t_half):
    if t_half and t_half > 0:
        return max(math.ceil((t_half * 5) / 24), 7)
    return 10

def calculate_sampling_schedule(t_half, tmax=None):
    if not t_half or t_half <= 0:
        t_half = 24
    points = [0]
    if tmax and tmax > 0:
        for h in [0.25, 0.5, 0.75, 1, 1.5, 2]:
            if h < tmax:
                points.append(h)
        points.extend([tmax * 0.8, tmax, tmax * 1.2])
    else:
        for h in [0.25, 0.5, 0.75, 1, 1.5, 2, 3, 4, 6, 8, 12]:
            points.append(h)
    cur = 24
    max_h = int(t_half * 7)
    while cur <= max_h:
        points.append(cur)
        cur += 12 if cur < 48 else 24
    if points[-1] < max_h:
        points.append(max_h)
    return sorted(set(round(x, 2) for x in points))

def choose_design(cv_intra, t_half, design_choice="автоматически", use_rsabe=False, regulator="EMA (Европа)", cv_threshold=30):
    if design_choice != "автоматически":
        if design_choice == "2x2":
            return "2x2", "Стандартный 2-периодный перекрёстный дизайн"
        elif design_choice == "2x2x3":
            return "2x2x3", "3-периодный перекрёстный дизайн"
        elif design_choice == "2x2x4":
            return "2x2x4", "4-периодный перекрёстный дизайн"
        elif design_choice == "parallel":
            return "parallel", "Параллельный дизайн"
    if t_half and t_half > 30:
        return "parallel", "Параллельный дизайн"
    if cv_intra is None:
        return "2x2", "Стандартный 2-периодный перекрёстный"
    cv_percent = cv_intra * 100 if cv_intra else 0
    if study_type == "однофазное (параллельное)":
        return "parallel", "Параллельный дизайн"
    if use_rsabe and cv_percent >= cv_threshold:
        if regulator == "FDA (США)":
            return "2x2x3", f"Репликативный дизайн с RSABE (FDA, CV={cv_percent:.1f}%)"
        else:  
            if cv_percent >= 50:
                return "2x2x4", f"4-периодный дизайн с RSABE (EMA, CV={cv_percent:.1f}%)"
            else:
                return "2x2x3", f"3-периодный дизайн с RSABE (EMA, CV={cv_percent:.1f}%)"
    if cv_percent <= 30:
        return "2x2", "2-периодный"
    elif cv_percent <= 50:
        return "2x2x3", "3-периодный"
    else:
        return "2x2x4", "4-периодный"

def calculate_sample_size(design, cv_intra, target_power, alpha, theta0=0.95, use_rsabe=False, regulator="EMA (Европа)"):
    try:
        if cv_intra is None or cv_intra <= 0:
            cv_intra = 0.25
        z_alpha = {0.1: 1.282, 0.05: 1.645, 0.025: 1.96, 0.01: 2.326}.get(alpha, 1.645)
        z_beta = {0.7: 0.525, 0.75: 0.675, 0.8: 0.84, 0.85: 1.04, 0.9: 1.28, 0.95: 1.645}.get(round(target_power, 2), 0.84)
        if use_rsabe and cv_intra > 0.3:
            if regulator == "FDA (США)":
                theta1 = max(0.8, math.exp(-0.893 * cv_intra))
                theta2 = min(1.25, math.exp(0.893 * cv_intra))
            else:
                scaled_limit = math.exp(0.76 * cv_intra)
                theta1 = max(0.8, min(0.6984, 1/scaled_limit))
                theta2 = min(1.25, max(1.4319, scaled_limit))
        else:
            theta1, theta2 = 0.80, 1.25      
        ln_t0, ln_t1, ln_t2 = math.log(theta0), math.log(theta1), math.log(theta2)
        delta = min(abs(ln_t0 - ln_t1), abs(ln_t0 - ln_t2))
        n_base = ((z_alpha + z_beta)** 2 * cv_intra** 2) / (delta ** 2)
        mult = {"2x2": 2.0, "2x2x3": 3.0, "2x2x4": 4.0, "parallel": 4.0}.get(design, 2.0)
        n_raw = int(math.ceil(n_base * mult))
        if design == "parallel":
            min_n = 20
        else:
            min_n = 12
        n_raw = max(n_raw, min_n)
        if design != "parallel" and n_raw % 2:
            n_raw += 1  
        return n_raw, target_power
    except Exception as e:
        st.warning(f"Ошибка расчёта выборки: {e}")
        if design == "parallel":
            n_approx = int(40 * (cv_intra ** 2) * (target_power / 0.8) * 100)
            n_approx = max(n_approx, 20)
        else:
            n_approx = int(24 * (cv_intra ** 2) * (target_power / 0.8) * 100)
            n_approx = max(n_approx, 12)
            if n_approx % 2:
                n_approx += 1   
        return n_approx, target_power

def create_word_document(synopsis_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    lines = synopsis_text.split('\n')
    for line in lines:
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('=') and 'СИНОПСИС' in line:
            continue  
        elif line.startswith('=' * 50) or line.startswith('-' * 50):
            doc.add_paragraph('_' * 50)
        elif line and len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', '.\t']:
            doc.add_heading(line, level=1)
        elif line.startswith('  ') or line.startswith('\t'):
            p = doc.add_paragraph(line.lstrip())
            p.paragraph_format.left_indent = Inches(0.25)
        elif line and len(line) > 2 and line[0].isdigit() and line[1] == '.':
            p = doc.add_paragraph(style='List Number')
            p.text = line
        else:
            doc.add_paragraph(line)
    
    return doc

def generate_detailed_synopsis(params, studies):
    times = params['sampling_schedule']
    if len(times) <= 8:
        points_str = ", ".join(str(t) for t in times)
    else:
        points_str = ", ".join(str(t) for t in times[:8]) + f" и далее до {times[-1]} ч"

    total_blood_samples = len(times) * 2
    total_blood_volume_pk = total_blood_samples * 5
    system_fill = total_blood_samples * 0.5
    total_blood_volume_all = total_blood_volume_pk + system_fill + 50
    total_plasma_samples = len(times) * 2 * params['n_final']

    day1_dose = 1
    day1_end = params['period_duration']
    day2_start = params['washout_days'] + 1
    day2_dose = params['washout_days'] + 1
    day2_end = params['washout_days'] + params['period_duration']
    day_fu = params['washout_days'] + params['period_duration'] + 7

    tmax_double = params.get('tmax', 4.0) * 2

    def safe_format(v, fmt, default="—"):
        if v is None:
            return default
        try:
            return fmt.format(v)
        except:
            return str(v)

    cv = safe_format(params['cv'] * 100 if params['cv'] else None, "{:.1f}")
    thalf = safe_format(params['t_half'], "{:.1f}")
    tmax = safe_format(params.get('tmax'), "{:.1f}")
    gmr = safe_format(params['expected_gmr'], "{:.2f}")
    n_raw = safe_format(params['n_raw'], "{}")
    n_final = safe_format(params['n_final'], "{}")
    n_group = safe_format(params['n_per_group'], "{}")
    n_screen = safe_format(params.get('n_screening', '—'), "{}")
    power = safe_format(params['achieved_power'] * 100 if params['achieved_power'] else None, "{:.1f}")
    dropout = safe_format(params['dropout_rate'], "{}")
    target_pow = safe_format(params['target_power'] * 100 if params['target_power'] else None, "{:.1f}")
    alpha_fmt = safe_format(params['alpha'], "{}")
    wash = safe_format(params['washout_days'], "{}")
    period_dur = safe_format(params['period_duration'], "{}")
    total_dur = safe_format(params['total_duration'], "{}")
    age_min_f = safe_format(params['age_min'], "{}")
    age_max_f = safe_format(params['age_max'], "{}")
    bmi_min_f = safe_format(params['bmi_min'], "{:.1f}")
    bmi_max_f = safe_format(params['bmi_max'], "{:.1f}")
    w_min_m = safe_format(params['weight_min_m'], "{:.1f}")
    w_min_f = safe_format(params['weight_min_f'], "{:.1f}")
    w_max = safe_format(params['weight_max'], "{:.1f}")

    regimen_text = "натощак" if params['regimen'] == "натощак" else "после еды"

    synopsis = f"""
    
    СИНОПСИС ПРОТОКОЛА КЛИНИЧЕСКОГО ИССЛЕДОВАНИЯ
    
    1. ИДЕНТИФИКАЦИОННАЯ ИНФОРМАЦИЯ
    Спонсор:                                    {params['sponsor']}
    Номер протокола:                            {params['study_number']}
    Версия протокола:                           {params['protocol_version']}
    Дата протокола:                             {params['protocol_date']}
    Исследовательский центр:                    {params['research_center']}
    Биоаналитическая лаборатория:                {params['bioanalytical_lab']}
    
    2. НАЗВАНИЕ ИССЛЕДОВАНИЯ
    Открытое рандомизированное перекрестное исследование сравнительной фармакокинетики 
    и биоэквивалентности препаратов {params['test_product']}, {params['dosage_form']} {params['dose']} 
    (производитель: {params['manufacturer_t']}) и {params['reference_product']}, {params['dosage_form']} {params['dose']} 
    (производитель: {params['manufacturer_r']}, регистрационное удостоверение: {params['registration_number_r']}) 
    у здоровых добровольцев.
    
    3. ЦЕЛИ И ЗАДАЧИ ИССЛЕДОВАНИЯ
    Основная цель:
    Оценка сравнительной фармакокинетики и биоэквивалентности препаратов {params['test_product']} 
    и {params['reference_product']} {regimen_text} у здоровых добровольцев.
    
    Дополнительная цель:
    Сравнительная оценка безопасности однократного приема препаратов {params['test_product']} 
    у здоровых добровольцев.
    
    Задачи исследования:
    1. Определить концентрацию {params['inn']} в плазме крови добровольцев после однократного 
       применения сравниваемых препаратов.
    2. Оценить фармакокинетические параметры (Cmax, AUC0-t, AUC0-∞, Tmax, T½, Kel) 
       и относительную биодоступность.
    3. Провести статистический анализ для подтверждения биоэквивалентности.
    4. Оценить профиль безопасности (частоту НЯ/СНЯ, изменения лабораторных показателей, 
       физикального осмотра, ЭКГ, жизненно важных функций).
    
    4. ДИЗАЙН ИССЛЕДОВАНИЯ
    Согласно Правилам ЕАЭС (Решение №85), для сравнительной оценки фармакокинетики 
    стандартным дизайном является открытое рандомизированное двухпериодное перекрестное 
    исследование в двух группах с однократным приемом дозы. Периоды приема разделяются 
    отмывочным периодом продолжительностью не менее 5 периодов полувыведения (T½ = {thalf} ч).
    
    Выбран дизайн: {params['design']} – {params['design_description']}."""
    if params.get('study_type') == "однофазное (параллельное)":
        synopsis += f"\nВыбран однофазный параллельный дизайн, так как {params.get('design_description', '')}.\n"
    else:
        synopsis += f"\nВыбран двухфазный перекрестный дизайн, стандартный для исследований биоэквивалентности.\n"
    
    if params.get('use_rsabe'):
        synopsis += f"\nПрименяется метод RSABE (Reference-Scaled Average Bioequivalence) согласно требованиям {params.get('regulator', 'EMA')} для высоковариабельных препаратов (CVintra = {cv}%).\n"
    synopsis += f"""
    5. МЕТОДОЛОГИЯ ИССЛЕДОВАНИЯ
    Исследование проводится с участием здоровых добровольцев, соответствующих критериям 
    включения/невключения и подписавших информированное согласие.
    
    Исследование состоит из следующих периодов:
    - Период скрининга (до 14 дней)
    - Два периода ФК исследования (Период 1 и Период 2) по {period_dur} дня каждый
    - Отмывочный период ({wash} дней)
    - Период последующего наблюдения (7 дней)
    
    Рандомизация и группы:
    Добровольцы распределяются в две группы в соотношении 1:1 методом блочной рандомизации:
      Группа 1 (n={n_group}): последовательность T / R
      Группа 2 (n={n_group}): последовательность R / T
    
    Период скрининга:
    После подписания информированного согласия добровольцы проходят процедуры скрининга 
    (физикальное обследование, лабораторные тесты, ЭКГ, сбор анамнеза) для оценки соответствия 
    критериям. Дата подписания согласия считается датой включения в исследование.
    
    Периоды ФК исследования:
    Добровольцы госпитализируются в центр вечером накануне приема препарата. Утром в день 
    приема (День {day1_dose} для Периода 1, День {day2_dose} для Периода 2) они получают 
    однократную дозу препарата {regimen_text} (не менее чем через 8 часов после последнего 
    приема пищи), запивая 200 мл воды. Добровольцы остаются в центре не менее 24 часов после 
    дозирования для отбора проб крови и оценки безопасности.
    
    График отбора проб крови на ФК:
    За 30 минут до приема и через {points_str} после приема. Всего {len(times)} точек в каждом 
    периоде. Объем одной пробы – 5 мл.
    
    Общий объем крови на ФК для одного добровольца за два периода:
    {total_blood_volume_pk} мл ({total_blood_samples} проб по 5 мл + заполнение системы 
    {system_fill:.1f} мл). Дополнительно 50 мл забирается для клинических анализов. 
    Итого {total_blood_volume_all:.0f} мл за всё исследование.
    
    Всего образцов плазмы, которые будут переданы в лабораторию:
    {total_plasma_samples} ({len(times)} точек × 2 периода × {params['n_final']} добровольцев).
    
    Отмывочный период:
    Длительность {wash} дней (5 × T½ = {5 * params['t_half']:.1f} ч) обеспечивает полное выведение 
    препарата перед вторым периодом. Во время отмывки проводится оценка безопасности.
    
    Период последующего наблюдения:
    Визит на {day_fu} день после последнего приема препарата для оценки НЯ/СНЯ.
    
    6. КОЛИЧЕСТВО ДОБРОВОЛЬЦЕВ
    Расчет размера выборки выполнен методом Power TOST (пакет PowerTOST в R) на основе 
    следующих параметров:
    - Ожидаемый внутрииндивидуальный коэффициент вариации (CVintra): {cv}%
    - Ожидаемое отношение средних геометрических (GMR): {gmr}
    - Целевая мощность: {target_pow}%
    - Уровень значимости α = {alpha_fmt}
    
    Базовый размер выборки (без учета отсева): {n_raw} добровольцев.
    С учетом ожидаемого отсева {dropout}% итоговое число рандомизированных добровольцев: {n_final}.
    С учетом отсева на скрининге {params.get('screening_fail_rate', 20)}% планируется 
    скринировать до {n_screen} добровольцев.
    Добровольцы, досрочно завершившие исследование, не заменяются.
    
    7. КРИТЕРИИ ВКЛЮЧЕНИЯ
    Для включения в исследование добровольцы должны соответствовать всем следующим критериям:
    1. Подписанное информированное согласие до начала процедур скрининга.
    2. Мужчины и женщины в возрасте от {age_min_f} до {age_max_f} лет включительно.
    3. Верифицированный диагноз «здоров» по данным клинических, лабораторных и инструментальных 
       методов обследования.
    4. Индекс массы тела (ИМТ) от {bmi_min_f} до {bmi_max_f} кг/м², масса тела ≥{w_min_f} кг 
       для женщин, ≥{w_min_m} кг для мужчин и не более {w_max} кг.
    5. Артериальное давление: САД 100–129 мм рт.ст., ДАД 60–89 мм рт.ст.
    6. Частота сердечных сокращений 60–89 уд/мин.
    7. Частота дыхательных движений 12–20 в минуту.
    8. Температура тела 36,0–36,9°C.
    9. Отказ от алкоголя за 72 часа до скрининга и на время исследования.
    10. Согласие на использование надежных методов контрацепции в течение всего исследования 
        и 1 месяца после его окончания.
    
    8. КРИТЕРИИ НЕВКЛЮЧЕНИЯ
    Добровольцы не будут включены в исследование при наличии любого из следующих критериев:
    1. Острые или хронические заболевания сердечно-сосудистой, дыхательной, нервной, эндокринной, 
       опорно-двигательной, кроветворной, иммунной систем, почек, печени, ЖКТ, кожи, 
       онкологические заболевания.
    2. Хирургические вмешательства на ЖКТ в анамнезе (кроме аппендэктомии >1 года назад).
    3. Состояния, влияющие на абсорбцию, распределение, метаболизм или экскрецию препаратов.
    4. Острые инфекционные заболевания менее чем за 30 дней до скрининга.
    5. Прием любых лекарственных препаратов (включая витамины, БАД) менее чем за 30 дней до 
       скрининга.
    6. Донорство крови или плазмы (>450 мл) менее чем за 2 месяца до скрининга.
    7. Применение гормональных контрацептивов (для женщин) менее чем за 2 месяца до скрининга.
    8. Депо-инъекции или имплантация препаратов в течение 6 месяцев до скрининга.
    9. Отклонения от нормы в лабораторных или инструментальных показателях при скрининге.
    10. Положительный тест на алкоголь, наркотики, котинин.
    11. Положительный тест на беременность (для женщин).
    12. Положительные тесты на ВИЧ, гепатиты B и C, сифилис.
    13. Высокая вероятность проблем с венепункцией.
    14. Повышенная чувствительность к гепарину или тромбоцитопения в анамнезе.
    15. Участие в другом клиническом исследовании менее чем за 3 месяца.
    16. Гиперчувствительность к {params['inn']} или любому компоненту препаратов.
    17. Отягощенный аллергологический анамнез, лекарственная непереносимость.
    18. Непереносимость лактозы, дефицит лактазы, глюкозо-галактозная мальабсорбция.
    19. Злоупотребление алкоголем (>10 ед/неделю) или наркотиками.
    20. Дегидратация (диарея, рвота) за 24 часа до приема препарата.
    21. Планируемая госпитализация на время исследования.
    22. Нарушения сна, экстремальные физические нагрузки.
    23. Особая диета (вегетарианская, гипокалорийная и т.п.) за 30 дней до скрининга.
    24. Употребление продуктов, содержащих ксантин (кофе, чай, шоколад и др.) за 72 часа 
        до приема препарата, а также грейпфрута, помело, клюквы, зверобоя за 7 дней.
    25. Беременность, лактация, отказ от контрацепции.
    26. Незащищенный половой акт у женщин за 30 дней до скрининга.
    27. Неспособность соблюдать процедуры протокола по мнению исследователя.
    28. Противопоказания к применению {params['inn']} согласно инструкции.
    
    9. КРИТЕРИИ ИСКЛЮЧЕНИЯ 
    Доброволец может быть исключен из исследования по следующим причинам:
    1. Отзыв информированного согласия.
    2. Несоблюдение требований протокола (пропуск процедур, прием запрещенных препаратов, 
       нарушение диеты и т.д.).
    3. Включение с нарушением критериев.
    4. Возникновение ситуаций, угрожающих безопасности (аллергические реакции и т.п.).
    5. Развитие НЯ/СНЯ, требующих вывода.
    6. Необходимость лечения, влияющего на ФК параметры.
    7. Пропуск 2 и более проб крови подряд или 3 и более за период.
    8. Рвота или диарея в течение {tmax_double:.1f} часов после приема препарата 
       (удвоенное Tmax).
    9. Положительный тест на алкоголь или наркотики в ходе исследования.
    10. Положительный тест на беременность.
    11. Другие причины, препятствующие выполнению протокола.
    
    10. ИССЛЕДУЕМЫЙ ПРЕПАРАТ
    Наименование: {params['test_product']}
    Лекарственная форма: {params['dosage_form']}
    Дозировка: {params['dose']}
    Состав на одну единицу:
      Действующее вещество: {params['inn']}
      Вспомогательные вещества: {params['auxiliary_substances_t']}
    Схема приема: однократно {regimen_text} в День {day1_dose} (для группы T/R) или 
      в День {day2_dose} (для группы R/T).
    Условия хранения: {params['storage_conditions']}
    Производитель: {params['manufacturer_t']}
    
    11. РЕФЕРЕНТНЫЙ ПРЕПАРАТ
    Наименование: {params['reference_product']}, МНН: {params['inn']}
    Лекарственная форма: {params['dosage_form']}
    Дозировка: {params['dose']}
    Состав на одну единицу:
      Действующее вещество: {params['inn']}
      Вспомогательные вещества: {params['auxiliary_substances_r']}
    Схема приема: однократно {regimen_text} в День {day1_dose} (для группы R/T) или 
      в День {day2_dose} (для группы T/R).
    Условия хранения: {params['storage_conditions']}
    Производитель: {params['manufacturer_r']}
    Регистрационное удостоверение: {params['registration_number_r']}
    Выбор референтного препарата обоснован Правилами ЕАЭС (Решение №85): 
    {params['reference_product']} является оригинальным препаратом, зарегистрированным в РФ.
    
    12. ПЕРИОДЫ ИССЛЕДОВАНИЯ 
    Период скрининга (Визит 1): дни -14 … -1. Оценка соответствия критериям, подписание ИС.
    
    Период 1 ФК (Визит 2):
    - День 0: госпитализация вечером.
    - День {day1_dose}: утренний прием препарата, отбор проб согласно графику.
    - Дни {day1_dose} – {day1_end}: пребывание в центре, отбор проб, оценка безопасности.
    - День {day1_end}: выписка.
    
    Отмывочный период: дни {day1_end+1} – {params['washout_days']}. Добровольцы находятся дома, 
    соблюдают ограничения, ведут дневник.
    
    Период 2 ФК (Визит 3):
    - День {day2_start}: повторная госпитализация.
    - День {day2_dose}: утренний прием другого препарата, отбор проб.
    - Дни {day2_dose} – {day2_end}: пребывание в центре, отбор проб, оценка безопасности.
    - День {day2_end}: выписка.
    
    Период последующего наблюдения (Визит 4):
    - День {day_fu}: визит в центр для оценки НЯ/СНЯ, финальное обследование.
    
    Незапланированный визит: при необходимости по решению исследователя.
    Визит досрочного завершения: при выбытии добровольца.
    
    13. ПРОДОЛЖИТЕЛЬНОСТЬ УЧАСТИЯ
    Максимальная продолжительность участия одного добровольца: {total_dur} дней 
    (скрининг до 14 дней + 2 ФК-периода по {period_dur} дня + отмывка {wash} дней + 7 дней наблюдения).
    
    14. ИЗУЧАЕМЫЕ ФАРМАКОКИНЕТИЧЕСКИЕ ПАРАМЕТРЫ
    
    Первичные (для оценки биоэквивалентности):
    - Cmax – максимальная концентрация в плазме.
    - AUC0-t – площадь под кривой «концентрация-время» от 0 до последней измеренной точки.
    - AUC0-∞ – площадь под кривой, экстраполированная до бесконечности.
    
    Вторичные:
    - Tmax – время достижения Cmax.
    - T½ – период полувыведения.
    - Kel – константа скорости элиминации.
    
    15. АНАЛИТИЧЕСКИЙ МЕТОД
    
    Концентрацию {params['inn']} в плазме крови определяют методом высокоэффективной 
    жидкостной хроматографии с тандемным масс-спектрометрическим детектированием (ВЭЖХ-МС/МС). 
    Метод будет полностью валидирован в соответствии с требованиями ЕАЭС.
    """

    if params.get('use_rsabe'):
        synopsis += f"""
    16. КРИТЕРИИ БИОЭКВИВАЛЕНТНОСТИ (RSABE)
    Для высоковариабельного препарата (CVintra = {cv}%) применяется метод 
    масштабированных границ биоэквивалентности (RSABE) согласно требованиям {params.get('regulator', 'EMA')}.
    
    Препараты считаются биоэквивалентными, если:
    - Точечная оценка отношения средних геометрических (T/R) находится в пределах 80.00% – 125.00%
    - 90% доверительный интервал для масштабированного критерия находится в пределах 
      расширенных границ, рассчитанных с учетом внутрииндивидуальной вариабельности.
    """
    else:
        synopsis += f"""
    16. КРИТЕРИИ БИОЭКВИВАЛЕНТНОСТИ
    Препараты считаются биоэквивалентными, если 90% доверительные интервалы для отношений 
    средних геометрических (T/R) для Cmax и AUC0-t находятся в пределах 80,00% – 125,00% 
    (α = {alpha_fmt}).
    """
    synopsis += f"""
    17. АНАЛИЗ БЕЗОПАСНОСТИ
    
    Безопасность оценивается по:
    - Нежелательным явлениям (НЯ) и серьезным НЯ (СНЯ) – регистрация, классификация по MedDRA, 
      оценка связи с препаратом.
    - Динамике лабораторных показателей (клинический и биохимический анализы крови, общий анализ мочи).
    - Показателям ЭКГ в 12 отведениях.
    - Данным физикального осмотра.
    - Жизненным показателям (АД, ЧСС, ЧДД, температура тела).
    
    18. РАСЧЕТ РАЗМЕРА ВЫБОРКИ
    
    Расчет выполнен с использованием программного обеспечения R (пакет PowerTOST) 
    методом Power TOST для двухпериодного перекрестного дизайна. Исходные параметры: 
    CVintra = {cv}%, ожидаемое GMR = {gmr}, целевая мощность {target_pow}%, α = {alpha_fmt}. 
    Полученный базовый размер выборки {n_raw} добровольцев скорректирован с учетом 
    {dropout}% отсева до {n_final} рандомизированных добровольцев.
    
    19. МЕТОДЫ СТАТИСТИЧЕСКОГО АНАЛИЗА
    
    Первичный анализ биоэквивалентности проводится на логарифмически преобразованных 
    показателях Cmax и AUC с использованием дисперсионного анализа (ANOVA) с фиксированными 
    факторами: последовательность, период, препарат и случайным фактором «субъект в последовательности».
    Для Tmax используется непараметрический анализ (критерий Вилкоксона).
    Гипотезы:
    H₀₁: μT/μR ≤ 0.80 против H₁₁: μT/μR > 0.80
    H₀₂: μT/μR ≥ 1.25 против H₁₂: μT/μR < 1.25
    Описательная статистика рассчитывается для всех параметров. Анализ безопасности выполняется 
    методами описательной статистики.
    
    20. ЗАСЛЕПЛЕНИЕ И РАНДОМИЗАЦИЯ
    
    Исследование открытое, однако биоаналитическая лаборатория не имеет доступа к 
    рандомизационному коду до завершения анализа. Рандомизация блочная (размер блока 4) 
    без стратификации, соотношение 1:1, выполняется с использованием программы IWRS iRand.
    
    21. ЭТИЧЕСКИЕ И РЕГУЛЯТОРНЫЕ АСПЕКТЫ
    
    Исследование проводится в соответствии с:
    - Хельсинкской декларацией ВМА (последняя редакция).
    - Правилами надлежащей клинической практики ЕАЭС (Решение №79).
    - Правилами проведения исследований биоэквивалентности ЕАЭС (Решение №85).
    - Законодательством Российской Федерации.
    Страхование жизни и здоровья добровольцев осуществляется компанией 
    {params.get('insurance_company', 'ООО «Страховая Компания»')}.
    
    22. НОМЕР ВЕРСИИ ПРОТОКОЛА И ДАТА
    
    Версия {params['protocol_version']} от {params['protocol_date']}
    """

    if studies:
        sources = "\n\nИСПОЛЬЗОВАННЫЕ ИСТОЧНИКИ ДАННЫХ\n"
        sources += "-" * 80 + "\n"
        for i, s in enumerate(studies, 1):
            if not isinstance(s, dict):
                continue
            src = s.get('source', 'Источник')
            sources += f"\n{i}. {src}\n"
            if s.get('title'):
                sources += f"   Название: {s['title']}\n"
            if s.get('name'):
                sources += f"   Препарат: {s['name']}\n"
            if s.get('journal'):
                sources += f"   Журнал: {s['journal']} ({s.get('year','')})\n"
            if s.get('authors'):
                sources += f"   Авторы: {s['authors']}\n"
            if s.get('url'):
                sources += f"   Ссылка: {s['url']}\n"
            if s.get('extracted_params'):
                plist = [f"{k.upper()}={v}" for k, v in s['extracted_params'].items() if v]
                if plist:
                    sources += f"   Извлеченные параметры: {', '.join(plist)}\n"
        synopsis += sources
    else:
        synopsis += "\n\n*Данные из внешних источников не использовались.*\n"

    synopsis += f"\nДата генерации: {datetime.today().strftime('%d.%m.%Y %H:%M')}\n"
    syn_len = len(synopsis)
    synopsis += f"\nДлина документа: {syn_len} знаков\n"
    return synopsis

if run_button:
    if age_max < age_min:
        st.error("Максимальный возраст меньше минимального!")
        st.stop()
    if bmi_max < bmi_min:
        st.error("Максимальный ИМТ меньше минимального!")
        st.stop()
    design_choice = design_choice if 'design_choice' in dir() else "автоматически"
    study_type = study_type if 'study_type' in dir() else "двухфазное"
    use_rsabe = use_rsabe if 'use_rsabe' in dir() else False
    regulator = regulator if 'regulator' in dir() else "EMA (Европа)"
    cv_threshold = cv_threshold if 'cv_threshold' in dir() else 30
    
    with st.spinner("Поиск и анализ данных..."):
        if use_auto_pk:
            pk_data = fetch_pk_data_all(inn, use_pubmed, use_grls, use_drugbank, use_pkdb)
            cv = pk_data.get("cv_intra") if pk_data.get("cv_intra") is not None else manual_cv
            cv_original = cv 
            cv = cv / 100     
            t_half = pk_data.get("t_half") if pk_data.get("t_half") is not None else manual_thalf
            tmax = pk_data.get("tmax") if pk_data.get("tmax") is not None else tmax_hours
            studies = pk_data.get("studies", [])
            found_params = []
            if pk_data.get("cv_intra") is not None:
                found_params.append(f"CV={pk_data['cv_intra']:.1f}%")
            if pk_data.get("t_half") is not None:
                found_params.append(f"T½={pk_data['t_half']:.1f} ч")
            if pk_data.get("tmax") is not None:
                found_params.append(f"Tmax={pk_data['tmax']:.1f} ч")
            if found_params:
                st.success(f"Найдены параметры: {', '.join(found_params)}")
            else:
                st.info(f"Используется ручной ввод: CV={manual_cv}%, T½={manual_thalf} ч, Tmax={tmax_hours} ч")
            if studies:
                with st.expander(f"Найдено исследований: {len(studies)}", key="studies_expander"):
                    for i, s in enumerate(studies, 1):
                        if not isinstance(s, dict):
                            st.write(f"{i}. {s} (некорректный формат)")
                            continue
                        st.markdown(f"**{i}. {s.get('title', s.get('name', 'Без названия'))}**")
                        if s.get('journal'):
                            st.write(f" {s['journal']} ({s.get('year','')})")
                        if s.get('authors'):
                            st.write(f" {s['authors']}")
                        if s.get('url'):
                            st.markdown(f" [Ссылка]({s['url']})")
                        if s.get('instruction_url'):
                            st.markdown(f" [Инструкция]({s['instruction_url']})")
                        if s.get('extracted_params'):
                            pl = [f"{k.upper()}={v}" for k, v in s['extracted_params'].items() if v]
                            if pl:
                                st.write(f" Параметры: {', '.join(pl)}")
                        if s.get('abstract'):
                            with st.expander("Аннотация", key=f"abstract_{i}_{s.get('pmid', i)}"):
                                st.write(s['abstract'])
                        st.write("---")
        else:
            cv = manual_cv / 100
            t_half = manual_thalf
            tmax = tmax_hours
            studies = []
            cv_original = manual_cv

        design, design_desc = choose_design(cv, t_half, design_choice, use_rsabe, regulator, cv_threshold)
        st.info(f"Дизайн: {design} – {design_desc}")
        n_raw, achieved_power = calculate_sample_size(design, cv, target_power, alpha, expected_gmr, use_rsabe, regulator)
        n_final = int(math.ceil(n_raw * (1 + dropout_rate / 100)))
        n_per_group = n_final // 2
        n_screening = int(math.ceil(n_final / (1 - screening_fail_rate / 100)))

        st.success(f" **Размер выборки:** {n_raw} → {n_final} (отсев {dropout_rate}%)")

        washout_days = calculate_washout_period(t_half)
        sampling_schedule = calculate_sampling_schedule(t_half, tmax)
        period_duration = 4
        total_duration = 14 + period_duration + washout_days + period_duration + 7

        total_blood_samples = len(sampling_schedule) * 2
        total_blood_volume_pk = total_blood_samples * 5
        system_fill = total_blood_samples * 0.5
        total_blood_volume_all = total_blood_volume_pk + system_fill + 50

        params_dict = {
            'sponsor': sponsor, 'study_number': study_number, 'protocol_version': protocol_version,
            'protocol_date': protocol_date.strftime("%d.%m.%Y"), 'research_center': research_center,
            'bioanalytical_lab': bioanalytical_lab, 'test_product': test_product, 'reference_product': reference_product,
            'dosage_form': dosage_form, 'dose': dose, 'manufacturer_t': manufacturer_t, 'manufacturer_r': manufacturer_r,
            'registration_number_r': registration_number_r, 'auxiliary_substances_t': auxiliary_substances_t,
            'auxiliary_substances_r': auxiliary_substances_r, 'storage_conditions': storage_conditions, 'inn': inn,
            'regimen': regimen, 'gender': gender, 'age_min': age_min, 'age_max': age_max, 'bmi_min': bmi_min,
            'bmi_max': bmi_max, 'weight_min_m': weight_min_m, 'weight_min_f': weight_min_f, 'weight_max': weight_max,
            'cv': cv, 't_half': t_half, 'tmax': tmax, 'expected_gmr': expected_gmr, 'design': design,
            'design_description': design_desc, 'target_power': target_power, 'alpha': alpha, 'dropout_rate': dropout_rate,
            'n_raw': n_raw, 'n_final': n_final, 'n_per_group': n_per_group, 'n_screening': n_screening,
            'achieved_power': achieved_power, 'washout_days': washout_days, 'sampling_schedule': sampling_schedule,
            'period_duration': period_duration, 'total_duration': total_duration,
            'insurance_company': "ООО «Страховая Компания»",
             'design_choice': design_choice,
            'study_type': study_type,
            'use_rsabe': use_rsabe,
            'regulator': regulator if use_rsabe else "не применяется",
            'cv_threshold': cv_threshold
        }

        synopsis = generate_detailed_synopsis(params_dict, studies)
        if synopsis:
            st.markdown("---")
            st.header("Черновик синопсиса")
            st.info(f"Длина: {len(synopsis)} знаков")
            st.markdown(synopsis)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button("Скачать .txt", data=synopsis,
                                   file_name=f"Синопсис_{study_number}_{inn}_{datetime.today().strftime('%Y%m%d')}.txt",
                                   mime="text/plain"
                                   )
            with col2:
                if WORD_AVAILABLE:
                    doc = create_word_document(synopsis)
                    bio = BytesIO()
                    doc.save(bio)
                    st.download_button(
                        label="Скачать .docx",
                        data=bio.getvalue(),
                        file_name=f"Синопсис_{study_number}_{inn}_{datetime.today().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.button("Word недоступен", disabled=True, help="Установите: pip install python-docx")
            with col3:
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <title>Синопсис {study_number}</title>
                    <style>
                        body {{
                            font-family: 'Times New Roman', Times, serif;
                            font-size: 12pt;
                            margin: 1.5cm;
                            line-height: 1.3;
                            white-space: pre-wrap;
                        }}
                        hr {{
                            border: none;
                            border-top: 1px solid black;
                            margin: 10px 0;
                        }}
                    </style>
                </head>
                <body>
                    {synopsis.replace(chr(10), '<br>')}
                </body>
                </html>
                """
    
                st.download_button(
                    "Скачать .html",
                    data=html_content,
                    file_name=f"Синопсис_{study_number}_{inn}_{datetime.today().strftime('%Y%m%d')}.html",
                    mime="text/html"
                )

            with st.expander("Детальные расчёты"):
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.metric("CVintra", f"{cv_original:.1f}%")
                    st.metric("T½", f"{t_half:.1f} ч")
                    st.metric("Tmax", f"{tmax:.1f} ч")
                    st.metric("GMR", f"{expected_gmr}")
                with c2:
                    st.metric("Дизайн", design)
                    st.metric("Мощность", f"{achieved_power*100:.1f}%")
                    st.metric("α", f"{alpha}")
                    st.metric("Отмывка", f"{washout_days} дн")
                with c3:
                    st.metric("Выборка", n_raw)
                    st.metric("Выборка", n_final)
                    st.metric("На группу", n_per_group)
                st.subheader("График забора")
                st.write(f"**Точек:** {len(sampling_schedule)}")
                st.write(f"**Часы:** {sampling_schedule}")
                st.write(f"**Последняя:** {sampling_schedule[-1]} ч")
                st.write(f"**Объём крови:** {total_blood_volume_all:.0f} мл")

            with st.expander("Исходные параметры"):
                st.json(params_dict)
        else:
            st.error("Не удалось сгенерировать синопсис.")
